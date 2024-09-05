from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 파일에서 텍스트 읽기
with open('article_0.txt', 'r', encoding='utf-8') as file:
    content = file.read()

# 파일 내용을 '\n' 단위로 분리하여 첫 번째 줄은 Title로, 나머지는 Content로 처리
lines = content.split('\n')
title = lines[0].strip()  # 첫 번째 줄은 제목
content = '\n'.join(lines[1:]).strip()  # 나머지 줄은 내용으로 처리

# Word 문서 생성
doc = Document()

# Title을 문서의 첫 번째 부분에 큰 글씨로 추가
doc.add_heading(title, level=1)

# Content를 문서에 추가
paragraph = doc.add_paragraph(content)

# 스타일 설정: 폰트 Georgia, 크기 12, 행간 1.5
run = paragraph.runs[0]
run.font.name = 'Georgia'
run.font.size = Pt(12)

# 문단의 행간 설정 (1.5배)
p = paragraph._element
pPr = p.get_or_add_pPr()
spacing = OxmlElement('w:spacing')
spacing.set(qn('w:line'), '360')  # 1.5배 행간 (240 * 1.5 = 360)
pPr.append(spacing)

# 파일 저장
doc.save('RFK_Article_Styled.docx')

print("Word 파일이 생성되었습니다.")
